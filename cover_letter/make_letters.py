import argparse, csv, os, sys, re
from pathlib import Path
from docx2pdf import convert

# Optional deps: jinja2, python-docx
try:
    from jinja2 import Template
except ImportError:
    Template = None
try:
    from docx import Document
except ImportError:
    Document = None



COMPANY_RE  = re.compile(r"{{\s*company\s*}}")
POSITION_RE = re.compile(r"{{\s*position\s*}}")

def _merge_spans(spans):
    if not spans: return []
    spans = sorted(spans)
    merged = [spans[0]]
    for s, e in spans[1:]:
        ls, le = merged[-1]
        if s <= le:
            merged[-1] = (ls, max(le, e))
        else:
            merged.append((s, e))
    return merged

def _find_all_spans(text: str, phrase: str):
    spans = []
    start = 0
    while True:
        i = text.find(phrase, start)
        if i == -1: break
        spans.append((i, i + len(phrase)))
        start = i + len(phrase)
    return spans

def _first_sentence_span(text: str):
    """
    Very simple first-sentence detector: up to the first period, exclamation, or question mark
    followed by a space/newline/end. Handles commas etc. Fine for cover letters.
    """
    m = re.search(r'[.!?](?:\s|$)', text)
    if not m:
        # no terminal punctuation; bold the whole paragraph if short, else skip
        return (0, len(text)) if text.strip() else None
    return (0, m.end())  # include the punctuation

def _rebuild_paragraph_with_bold(paragraph, bold_spans):
    text = paragraph.text
    base_font_name = paragraph.runs[0].font.name if paragraph.runs else None
    base_font_size = paragraph.runs[0].font.size if paragraph.runs else None
    for i in range(len(paragraph.runs) - 1, -1, -1):
        paragraph._p.remove(paragraph.runs[i]._r)

    idx = 0
    for (s, e) in bold_spans:
        if idx < s:
            r = paragraph.add_run(text[idx:s]); r.bold = False
            if base_font_name: r.font.name = base_font_name
            if base_font_size: r.font.size = base_font_size
        r = paragraph.add_run(text[s:e]); r.bold = True
        if base_font_name: r.font.name = base_font_name
        if base_font_size: r.font.size = base_font_size
        idx = e
    if idx < len(text):
        r = paragraph.add_run(text[idx:]); r.bold = False
        if base_font_name: r.font.name = base_font_name
        if base_font_size: r.font.size = base_font_size


def bold_phrases_and_first_sentence(paragraph, phrases_to_bold, bold_first_sentence=True):
    text = paragraph.text
    if not text.strip():
        return

    spans = []
    # First sentence
    if bold_first_sentence:
        fs = _first_sentence_span(text)
        if fs:
            spans.append(fs)

    # Exact phrases
    for phrase in phrases_to_bold:
        if not phrase: continue
        spans.extend(_find_all_spans(text, phrase))

    spans = _merge_spans(spans)
    if spans:
        _rebuild_paragraph_with_bold(paragraph, spans)



def replace_in_paragraph_runs(paragraph, company_val: str, position_val: str):
    # Join the paragraph text once, regex-replace, then write back as a single run.
    original = paragraph.text
    new_text = COMPANY_RE.sub(company_val, original)
    new_text = POSITION_RE.sub(position_val, new_text)
    if new_text == original:
        return

    # preserve base font from first run if present
    base_font_name = paragraph.runs[0].font.name if paragraph.runs else None
    base_font_size = paragraph.runs[0].font.size if paragraph.runs else None

    for i in range(len(paragraph.runs) - 1, -1, -1):
        paragraph._p.remove(paragraph.runs[i]._r)
    r = paragraph.add_run(new_text)
    if base_font_name: r.font.name = base_font_name
    if base_font_size: r.font.size = base_font_size

def replace_in_docx(doc, context: dict):
    company  = context.get("company", "")
    position = context.get("position", "")
    for p in doc.paragraphs:
        replace_in_paragraph_runs(p, company, position)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph_runs(p, company, position)

def render_docx_template(template_path: Path, context: dict, out_docx: Path, bold_list=None):
    if Document is None:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    doc = Document(str(template_path))

    # 1) Replace placeholders everywhere
    context_local = {"company": context.get("company",""), "position": context.get("position","")}
    replace_in_docx(doc, context_local)

    # 2) Bold rules
    ALWAYS_BOLD = bold_list or []
    # Add the phrases you requested:
    ALWAYS_BOLD = list(set(ALWAYS_BOLD + [
        "University of Chicago",
        "Financial Markets Program",
        "dual major in Computer Science and Economics",
    ]))

    # Pass 1: paragraphs
    for p in doc.paragraphs:
        bold_phrases_and_first_sentence(p, ALWAYS_BOLD, bold_first_sentence=True)

    # Pass 2: tables (in case header lives in a table)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    bold_phrases_and_first_sentence(p, ALWAYS_BOLD, bold_first_sentence=True)

    doc.save(str(out_docx))


def to_pdf_with_libreoffice(input_path: Path, out_dir: Path):
    os.system(f'libreoffice --headless --convert-to pdf "{input_path}" --outdir "{out_dir}"')

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", required=True, help="Path to .docx or .txt template")
    ap.add_argument("--company", help="Company name (single run)")
    ap.add_argument("--position", default="", help="Position title (single run)")
    ap.add_argument("--csv", help="CSV with headers: company,position")
    ap.add_argument("--pdf", action="store_true", help="Also export PDF")
    ap.add_argument("--outdir", default="coverletters/out", help="Output directory")
    args = ap.parse_args()

    tpl = Path(args.template)
    outdir = Path(args.outdir); outdir.mkdir(parents=True, exist_ok=True)
    
    def generate_one(company: str, position: str = ""):
        safe = company.replace("/", "-").replace("\\", "-").strip()
        basename = f"Chris Low {safe} Cover Letter"
        context = {"company": company, "position": position}

        if tpl.suffix.lower() == ".docx":
            out_docx = outdir / f"{basename}.docx"
            render_docx_template(tpl, context, out_docx, bold_list=ALWAYS_BOLD)
            if args.pdf:
                out_pdf = outdir / f"{basename}.pdf"
                convert(str(out_docx), str(out_pdf))

        else:
            out_txt = outdir / f"{basename}.txt"
            out_txt.write_text(render_text_template(tpl, context), encoding="utf-8")

    if args.company:
        generate_one(args.company.strip(), args.position.strip())
    elif args.csv:
        with open(args.csv, newline="") as f:
            for row in csv.DictReader(f):
                company = (row.get("company") or "").strip()
                position = (row.get("position") or "").strip()
                if company:
                    generate_one(company, position)
    else:
        print("Provide --company COMPANY or --csv companies.csv")
        sys.exit(2)

if __name__ == "__main__":
    main()
